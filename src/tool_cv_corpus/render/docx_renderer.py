"""DOCX renderer.

``python-docx`` is an optional dependency (install extra ``[docx]``) so
systems that only want PDF or HTML do not carry a transitive XML stack.
We import inside the method and raise a clear install hint if it is
missing, rather than failing at module import and breaking the CLI's
``--help`` for users who never asked for DOCX.
"""

from __future__ import annotations

from pathlib import Path
from typing import ClassVar

from .base import RenderedResume

_INSTALL_HINT = (
    "DOCX output needs python-docx. Install with: uv pip install 'tool-cv-corpus[docx]'"
)


class DocxRenderer:
    """Render a RenderedResume as a simple .docx document."""

    name: ClassVar[str] = "docx"
    extensions: ClassVar[tuple[str, ...]] = (".docx",)

    def render(self, resume: RenderedResume, out_path: Path) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(_INSTALL_HINT) from exc

        target = out_path if out_path.suffix else out_path.with_suffix(".docx")
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(resume.person.full_name, level=0)
        if resume.headline or resume.person.headline:
            doc.add_paragraph(resume.headline or resume.person.headline or "")
        if resume.summary:
            doc.add_paragraph(resume.summary)

        if resume.roles:
            doc.add_heading("Experience", level=1)
            for role in resume.roles:
                period = f"{role.period.start} to {role.period.end or 'present'}"
                doc.add_paragraph(
                    f"{role.title}, {role.organization_id} ({period})",
                    style="List Bullet",
                )
                for ach in (a for a in resume.achievements if a.role_id == role.id):
                    doc.add_paragraph(ach.headline, style="List Bullet 2")

        if resume.skills:
            doc.add_heading("Skills", level=1)
            for skill in resume.skills:
                doc.add_paragraph(
                    f"{skill.name} ({skill.tier}, {skill.confidence})",
                    style="List Bullet",
                )

        doc.save(str(target))
        return target
